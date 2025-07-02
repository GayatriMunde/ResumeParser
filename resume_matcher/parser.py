import os
import re
from docx import Document

class ResumeParser:
    def parse_resume(self, filepath):
        document = Document(filepath)
        full_text = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())

        name = document.paragraphs[0].text.strip() if document.paragraphs else ""
        email = re.search(r"[\w\.-]+@[\w\.-]+", full_text)
        phone = re.search(r"\+?\d[\d\s\-\(\)]{7,}\d", full_text)

        education = self.extract_education(full_text)
        experience_years = self.extract_experience_years(full_text)
        skills = []  # Skip section-based detection for API use

        return {
            "name": name,
            "email": email.group(0) if email else None,
            "phone": phone.group(0) if phone else None,
            "education": education,
            "total_years_experience": experience_years,
            "text_blob": full_text
        }

    def parse_directory(self, path):
        resumes = {}
        for file in os.listdir(path):
            if file.endswith(".docx"):
                resumes[file] = self.parse_resume(os.path.join(path, file))
        return resumes

    def extract_education(self, text):
        degrees = ["bachelor", "master", "phd", "b.tech", "m.tech", "mba"]
        return [deg.title() for deg in degrees if deg in text.lower()]

    def extract_experience_years(self, text):
        matches = re.findall(r"(\d{4})\s*[-–]\s*(\d{4})", text)
        total = 0
        for start, end in matches:
            try:
                total += max(0, int(end) - int(start))
            except:
                continue
        return total

    def build_resume_blob(self, parsed_data):
        """
        Converts structured resume dict into a plain text blob.
        This helps in computing semantic similarity.
        """
        blob_parts = []

        for key, value in parsed_data.items():
            if isinstance(value, list):
                blob_parts.extend(value)
            elif isinstance(value, str):
                blob_parts.append(value)

        return " ".join(blob_parts)
