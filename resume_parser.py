import os
import re
from docx import Document


class ResumeParser:
    def __init__(self, skill_vocab=None):
        self.skill_vocab = [skill.lower() for skill in skill_vocab] if skill_vocab else None

    def parse_resume(self, filepath):
        document = Document(filepath)
        full_text = "\n".join([para.text.strip() for para in document.paragraphs if para.text.strip()])

        name = document.paragraphs[0].text.strip() if document.paragraphs else ""
        email = re.search(r"[\w\.-]+@[\w\.-]+", full_text)
        phone = re.search(r"\+?\d[\d\s\-\(\)]{7,}\d", full_text)

        skills = self.extract_skills(document, full_text)
        education = self.extract_education(full_text)
        experience_years = self.extract_experience_years(full_text)

        return {
            "name": name,
            "email": email.group(0) if email else None,
            "phone": phone.group(0) if phone else None,
            "skills": skills,
            "education": education,
            "total_years_experience": experience_years
        }

    def extract_skills(self, document, full_text):
        skills_section = False
        skills = []

        for para in document.paragraphs:
            text = para.text.strip().lower()
            if "skills" in text:
                skills_section = True
                continue
            if skills_section:
                if not text or any(kw in text for kw in ["experience", "education", "projects", "certifications"]):
                    break
                skills.extend(text.split(","))

        skills = [skill.strip() for skill in skills if skill.strip()]

        if not skills and self.skill_vocab:
            for skill in self.skill_vocab:
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, full_text.lower()):
                    skills.append(skill)

        return list(set(skills))

    def extract_education(self, text):
        degrees = ["bachelor", "master", "phd", "b.tech", "m.tech", "mba", "b.sc", "m.sc"]
        found = []
        for degree in degrees:
            if degree in text.lower():
                found.append(degree.title())
        return list(set(found))

    def extract_experience_years(self, text):
        patterns = [
            r"(\d{4})\s*[-–]\s*(\d{4})",
            r"from\s+(\d{4})\s+to\s+(\d{4})"
        ]
        total_years = 0
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            for start, end in matches:
                try:
                    total_years += max(0, int(end) - int(start))
                except:
                    continue
        return total_years

    def parse_directory(self, directory):
        resume_data = {}
        for file in os.listdir(directory):
            if file.lower().endswith('.docx'):
                path = os.path.join(directory, file)
                resume_data[file] = self.parse_resume(path)
        return resume_data

    def build_resume_blob(self, resume_data):
        lines = []

        if resume_data.get("name"):
            lines.append(resume_data["name"])

        if resume_data.get("skills"):
            lines.append("Skills: " + ", ".join(resume_data["skills"]))

        if resume_data.get("education"):
            lines.append("Education: " + ", ".join(resume_data["education"]))

        lines.append(f"Experience: {resume_data.get('total_years_experience', 0)} years")

        return " ".join(lines)
